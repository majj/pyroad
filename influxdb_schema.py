
# influxdb_schema.py
# metadata
# gen schema in yaml

### pip install influxdb
### pip install yaml
### pip install ruamel

### pip install python-docx


# from __future__ import print_function

import sys

from influxdb import InfluxDBClient
import pendulum
import ruamel
from ruamel.yaml import YAML
from ruamel.yaml.compat import StringIO


from docx import Document
from docx.shared import Inches







#document.add_picture('monty-truth.png', width=Inches(1.25))

### table = document.add_table(rows=1, cols=3)
### hdr_cells = table.rows[0].cells
### hdr_cells[0].text = 'Qty'
### hdr_cells[1].text = 'Id'
### hdr_cells[2].text = 'Desc'
### for item in recordset:
    ### row_cells = table.add_row().cells
    ### row_cells[0].text = str(item.qty)
    ### row_cells[1].text = str(item.id)
    ### row_cells[2].text = item.desc

### document.add_page_break()



HOST = "localhost"
PORT = 8086
USER = "root"
PASSWORD = "root"

class MyYAML(YAML):
    """"""
    
    def dump(self, data, stream=None, **kw):
        """ return string """
        
        inefficient = False
        if stream is None:
            inefficient = True
            stream = StringIO()
        YAML.dump(self, data, stream, **kw)
        if inefficient:
            return stream.getvalue()

class GenX(object):
    
    def __init__(self, db):
        self.db = db
        self.doc = Document()
        self.doc.add_heading("InfluxDB Schema", level=0)
        
        self.client = InfluxDBClient(HOST, PORT, USER, PASSWORD, db)

    def gen(self):

        yaml = MyYAML()

        self.measurements = []

        result = self.client.query(f'SHOW MEASUREMENTS ON "{self.db}";')
        
        for k in result.keys():
            for item in result[k]:
                self.measurements.append(item["name"])            
                
        for meas in self.measurements:
            print(f"{meas}...")
            
            
            self.doc.add_heading(f"Measurement:{meas}", level=1)
            
            self.doc.add_paragraph(f'schema for {meas}')
            
            table = self.doc.add_table(rows=1, cols=4)
            table.style='Medium Grid 1 Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Tags/Fields'          
            hdr_cells[1].text = 'Item'
            hdr_cells[2].text = 'DataType'
            hdr_cells[3].text = 'Desc'

            
            
            data = ruamel.yaml.comments.CommentedMap()#yaml.load(inp)
            
            data["measurement"] = meas
            
            data["time"] = pendulum.now().to_iso8601_string()
            data.yaml_add_eol_comment('gen by mabo', 'time')
            
            ### tags
            tags = {}
            result = self.client.query(f'SHOW TAG KEYS ON "{self.db}"  from "{meas}";')        
            for k in result.keys():
                for item in result[k]:
                    tags[item["tagKey"]] = "tag"  # 
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = "Tags"
                    row_cells[1].text = item["tagKey"]                    
                    row_cells[2].text = ""
                    row_cells[3].text = ""                    
                    
            if len(tags.keys())>0:
                data["tags"] = tags    
            
            ### fields
            fields = {}
            result = self.client.query(f'SHOW FIELD KEYS FROM "{self.db}"."autogen"."{meas}";')

            for k in result.keys():
                for item in result[k]:
                    # type & unit
                    fields[item["fieldKey"]] = {"type":item["fieldType"], "unit":"u"}
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = "Fields"
                    row_cells[1].text = item["fieldKey"]                    
                    row_cells[2].text = item["fieldType"]
                    row_cells[3].text = ""
                    
                    #print(item["fieldKey"], item["fieldType"])
            data["fields"] = fields
            
            yaml.indent(mapping=4, sequence=4, offset=0)
            
            f = open(f"{meas}.yaml", 'w')
            yaml.dump(data, f)
        
        
        t = pendulum.now().format('YYYYMMDDHHmmss')
        
        
        dt = pendulum.now().format('YYYY-MM-DD HH:mm:ss')
        self.doc.add_paragraph('\n\n')
        self.doc.add_paragraph(f'gen on {dt}')
        
        self.doc.save(f'influxdb-{t}.docx')
        print("done!")

def main():
    
    db = "eim"
    x = GenX(db)
    x.gen()

if __name__ == "__main__":
    
    main()
    
    
###
### SHOW DATABASES
### {"results":[{"statement_id":0,"series":[{"name":"databases","columns":["name"],"values":[["_internal"],["eim"]]}]}]}
###
###
### SHOW MEASUREMENTS ON "eim"
### {"results":[{"statement_id":0,"series":[{"name":"measurements","columns":["name"],"values":[["mts"],["new_mts"],["nt01"],["nt2"],["nt3"],["xyz"]]}]}]}
###
###
### SHOW FIELD KEYS FROM "eim"."autogen"."nt2"
###
### {"results":[{"statement_id":0,"series":[{"name":"nt2","columns":["fieldKey","fieldType"],"values":[["i","integer"],["j","string"],["k","float"]]}]}]}
###
### SHOW TAG KEYS ON "eim"  from "nt2"
### {"results":[{"statement_id":0,"series":[{"name":"nt2","columns":["tagKey"],"values":[["eqptno"]]}]}]}
###
### SHOW TAG VALUES ON "eim" WITH KEY = "eqptno"
###    